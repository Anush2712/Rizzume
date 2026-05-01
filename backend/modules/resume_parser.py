"""
Resume Parser: extracts plain text from PDF, DOCX, or TXT files.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class ResumeParser:
    def parse(self, file_path: Path) -> str:
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".txt":
            return file_path.read_text(encoding="utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _parse_pdf(self, path: Path) -> str:
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            logger.warning("pdfplumber not available, falling back to PyPDF2")
            return self._parse_pdf_fallback(path)

    def _parse_pdf_fallback(self, path: Path) -> str:
        try:
            import PyPDF2
            text_parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("Install pdfplumber or PyPDF2: pip install pdfplumber")

    def _parse_docx(self, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
