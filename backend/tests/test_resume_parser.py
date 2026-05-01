"""Tests for ResumeParser module."""

import io
import tempfile
from pathlib import Path

import pytest

from modules.resume_parser import ResumeParser


@pytest.fixture
def parser():
    return ResumeParser()


def make_txt(content: str) -> Path:
    f = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8")
    f.write(content)
    f.close()
    return Path(f.name)


class TestResumeParserTxt:
    def test_parse_plain_text(self, parser):
        path = make_txt("John Doe\nSoftware Engineer\nPython, AWS, Docker")
        result = parser.parse(path)
        assert "John Doe" in result
        assert "Python" in result
        path.unlink()

    def test_parse_empty_text(self, parser):
        path = make_txt("")
        result = parser.parse(path)
        assert result == ""
        path.unlink()

    def test_parse_unicode(self, parser):
        path = make_txt("Résumé – Ünïcödé tëxt 日本語")
        result = parser.parse(path)
        assert "Résumé" in result
        path.unlink()

    def test_unsupported_extension(self, parser):
        path = make_txt("content")
        bad_path = path.with_suffix(".xyz")
        path.rename(bad_path)
        with pytest.raises(ValueError, match="Unsupported"):
            parser.parse(bad_path)
        bad_path.unlink()


class TestResumeParserDocx:
    def test_parse_docx(self, parser):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_heading("Jane Smith", level=1)
        doc.add_paragraph("Senior Data Engineer")
        doc.add_paragraph("Skills: Spark, Kafka, Python")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        doc.save(str(path))

        result = parser.parse(path)
        assert "Jane Smith" in result
        assert "Spark" in result
        path.unlink()

    def test_parse_docx_with_table(self, parser):
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Python"
        table.cell(0, 1).text = "5 years"
        table.cell(1, 0).text = "AWS"
        table.cell(1, 1).text = "3 years"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        doc.save(str(path))

        result = parser.parse(path)
        assert "Python" in result
        path.unlink()
